"""Generate the downloadable Wenzhou 2026 competition rules as DOCX.

The rules document JSON is read from standard input so the public page copy is
the single content source. Run this script with the bundled document Python
runtime (python-docx is required).
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BODY_FONT = "Calibri"
BODY_EAST_ASIA_FONT = "宋体"
HEADING_EAST_ASIA_FONT = "黑体"
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
MUTED = RGBColor(0x66, 0x66, 0x66)


def set_style_font(style, *, ascii_font: str, east_asia_font: str, size: float) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def set_run_font(run, *, size: float, east_asia_font: str, bold: bool = False,
                 color: RGBColor | None = None) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)


def add_body_paragraph(document: Document, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.31)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, east_asia_font=BODY_EAST_ASIA_FONT)


def add_chapter_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=16,
        east_asia_font=HEADING_EAST_ASIA_FONT,
        bold=True,
        color=HEADING_BLUE,
    )


def build_document(rules: dict, output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    set_style_font(
        normal,
        ascii_font=BODY_FONT,
        east_asia_font=BODY_EAST_ASIA_FONT,
        size=11,
    )
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading = document.styles["Heading 1"]
    set_style_font(
        heading,
        ascii_font=BODY_FONT,
        east_asia_font=HEADING_EAST_ASIA_FONT,
        size=16,
    )
    heading.font.bold = True
    heading.font.color.rgb = HEADING_BLUE
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run(rules["title"])
    set_run_font(
        title_run,
        size=22,
        east_asia_font=HEADING_EAST_ASIA_FONT,
        bold=True,
    )

    for chapter in rules["chapters"]:
        add_chapter_heading(document, chapter["heading"])
        for paragraph_text in chapter.get("paragraphs", []):
            no_indent = paragraph_text.startswith(("主办单位：", "承办单位：", "协办单位：", "支持单位："))
            add_body_paragraph(document, paragraph_text, indent=not no_indent)
        for item_text in chapter.get("items", []):
            add_body_paragraph(document, item_text, indent=False)
        for signature_text in chapter.get("signature", []):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(signature_text)
            set_run_font(run, size=11, east_asia_font=BODY_EAST_ASIA_FONT)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    footer_run = footer.add_run("温州市青少年航空航天创新比赛组委会")
    set_run_font(footer_run, size=9, east_asia_font=BODY_EAST_ASIA_FONT, color=MUTED)

    document.core_properties.title = rules["title"]
    document.core_properties.subject = "2026年温州市青少年航空航天创新比赛大赛章程"
    document.core_properties.author = "温州市青少年航空航天创新比赛组委会"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("usage: generate-rules-document.py OUTPUT.docx")
    rules = json.loads(sys.stdin.buffer.read().decode("utf-8"))
    build_document(rules, Path(sys.argv[1]).resolve())


if __name__ == "__main__":
    main()
